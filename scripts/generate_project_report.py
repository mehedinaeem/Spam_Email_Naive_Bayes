from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import matplotlib.pyplot as plt

ROOT_DIR = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT_DIR / "reports"
SCRIPTS_DIR = ROOT_DIR / "scripts"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

# Metrics from current model runs
metrics = {
    "Naive Bayes": {
        "Train": 0.9970,
        "Validation": 0.9930,
        "Test": 0.9843,
    },
    "Decision Tree": {
        "Train": 1.0000,
        "Validation": 0.9572,
        "Test": 0.9616,
    },
    "Random Forest": {
        "Train": 1.0000,
        "Validation": 0.9852,
        "Test": 0.9773,
    },
}

# Create accuracy comparison plot
accuracy_plot = REPORTS_DIR / "accuracy_comparison.png"
labels = ["Train", "Validation", "Test"]
plt.figure(figsize=(8, 5))
for model_name, values in metrics.items():
    plt.plot(labels, [values[label] for label in labels], marker="o", label=model_name)
plt.title("Model Accuracy Comparison")
plt.ylabel("Accuracy")
plt.ylim(0.90, 1.01)
plt.grid(True, linestyle="--", alpha=0.5)
plt.legend()
plt.tight_layout()
plt.savefig(accuracy_plot)
plt.close()

# Helper to add monospace code block

def add_code_block(document, code_text):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    run = paragraph.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


doc = Document()
doc.add_heading("Spam Email Detection Project Report", level=0)
doc.add_paragraph(
    "This document summarizes the spam email detection project, including dataset details, model training code, evaluation results, and an accuracy comparison graph."
)

doc.add_heading("1. Project Overview", level=1)
doc.add_paragraph(
    "The project detects spam emails using text classification. The dataset contains email text and binary labels: 1 for spam and 0 for not spam."
)

doc.add_paragraph("The processed dataset is split into training, validation, and test sets.")

doc.add_heading("2. Dataset and Preprocessing", level=1)
doc.add_paragraph(
    "The text data is vectorized using sklearn's CountVectorizer with English stop words removed. The vectorizer is fitted only on the training split and reused for validation and test splits."
)

doc.add_heading("3. Models Trained", level=1)
doc.add_paragraph("Three classification models were trained:")
for model_name in metrics.keys():
    doc.add_paragraph(f"- {model_name}", style="List Bullet")

doc.add_heading("4. Evaluation Results", level=1)
table = doc.add_table(rows=1, cols=4)
header_cells = table.rows[0].cells
header_cells[0].text = "Model"
header_cells[1].text = "Train Accuracy"
header_cells[2].text = "Validation Accuracy"
header_cells[3].text = "Test Accuracy"
for model_name, values in metrics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = model_name
    row_cells[1].text = f"{values['Train']:.4f}"
    row_cells[2].text = f"{values['Validation']:.4f}"
    row_cells[3].text = f"{values['Test']:.4f}"

doc.add_paragraph(
    "The table above shows the accuracy of each trained model on train, validation, and test splits."
)

doc.add_heading("5. Accuracy Comparison Graph", level=1)
doc.add_paragraph(
    "The following graph compares the accuracy of all three models across train, validation, and test datasets."
)
doc.add_picture(str(accuracy_plot), width=Inches(6))

doc.add_heading("6. Saved Outputs", level=1)
doc.add_paragraph("Model files saved in the models/ directory:")
doc.add_paragraph("- models/naive_bayes_model.pkl", style="List Bullet")
doc.add_paragraph("- models/decision_tree_model.pkl", style="List Bullet")
doc.add_paragraph("- models/random_forest_model.pkl", style="List Bullet")
doc.add_paragraph("- models/vectorizer.pkl", style="List Bullet")

doc.add_paragraph("Report files saved in the reports/ directory:")
doc.add_paragraph("- reports/classification_report.txt", style="List Bullet")
doc.add_paragraph("- reports/decision_tree_classification_report.txt", style="List Bullet")
doc.add_paragraph("- reports/random_forest_classification_report.txt", style="List Bullet")
doc.add_paragraph("- reports/project_report.docx", style="List Bullet")

doc.add_heading("7. Code Summary", level=1)
doc.add_paragraph(
    "The main training logic is the same for all three models: load data, vectorize text, fit the classifier, evaluate splits, save the model, and save the report."
)

# Include the core training function code from each script
for script_path in [SCRIPTS_DIR / "train_model.py", SCRIPTS_DIR / "train_decision_tree.py", SCRIPTS_DIR / "train_random_forest.py"]:
    doc.add_heading(f"Code: {script_path.name}", level=2)
    code_text = script_path.read_text(encoding="utf-8")
    add_code_block(doc, code_text)

doc.add_heading("8. Commands to Run", level=1)
doc.add_paragraph("Use the following commands from the project root:")
doc.add_paragraph("python scripts/train_model.py", style="List Bullet")
doc.add_paragraph("python scripts/train_decision_tree.py", style="List Bullet")
doc.add_paragraph("python scripts/train_random_forest.py", style="List Bullet")
doc.add_paragraph("python scripts/predict.py", style="List Bullet")

doc.save(REPORTS_DIR / "project_report.docx")
print(f"Report generated at: {REPORTS_DIR / 'project_report.docx'}")
print(f"Accuracy graph generated at: {accuracy_plot}")
